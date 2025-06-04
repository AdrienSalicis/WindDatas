import pandas as pd
from meteostat import Stations, Daily
from datetime import datetime
import tkinter as tk
from tkinter import simpledialog
from docx import Document
import os

# Charger les sites à partir du CSV
def charger_sites(csv_path):
    df = pd.read_csv(csv_path, sep=';')  # Assurez-vous que le CSV utilise ';' comme séparateur
    df['latitude'] = df['latitude'].str.replace(',', '.').astype(float)  # Convertir les coordonnées en float
    df['longitude'] = df['longitude'].str.replace(',', '.').astype(float)  # Convertir les coordonnées en float
    return df

# Fonction pour choisir la station à partir de Meteostat
def choisir_station_meteostat(stations):
    # Convertir les stations en DataFrame pour pouvoir itérer dessus
    stations_df = stations.fetch()
    options = []
    
    for i, row in stations_df.iterrows():
        try:
            # Assurer que les valeurs sont des chaînes de caractères
            name = str(row['name'])
            country = str(row['country'])
            # Créer l'option sans inclure la distance
            options.append(f"{str(i+1)}. {name} ({country})")
        except ValueError as e:
            print(f"Erreur de conversion pour la station {row['name']} avec la distance {row['distance']}: {e}")
            continue
    
    # Créer une interface tkinter pour choisir la station
    root = tk.Tk()
    root.withdraw()  # Cacher la fenêtre principale
    selection = simpledialog.askinteger("Choisir une station", f"Sélectionnez une station :\n" + "\n".join(options))

    if selection is not None and 1 <= selection <= len(stations_df):
        return stations_df.iloc[selection - 1]
    else:
        print("Sélection invalide")
        return None


# Récupérer les données de vent pour une station Meteostat donnée
def get_meteostat_wind_data(lat, lon, start, end):
    stations = Stations()
    stations = stations.nearby(lat, lon)
    
    # Choisir la station la plus proche
    station = choisir_station_meteostat(stations)
    
    if station is not None:
        # Récupérer les données météorologiques pour la station choisie
        data = Daily(station['id'], start, end)
        data = data.fetch()
        
        # Extraire les données relatives au vent
        wind_data = data[['wsp', 'gst', 'wdd', 'temp', 'pres', 'dew', 'prcp', 'rh']]

        # Ajouter les données maximales de vent (si disponibles)
        wind_data['wsp_max'] = data['wsp'].max()
        wind_data['gst_max'] = data['gst'].max()
        
        return wind_data, station
    else:
        print("Aucune station sélectionnée.")
        return None, None

# Générer un rapport Word organisé par pays
def generate_report(sites_data, csv_output_path):
    doc = Document()
    doc.add_heading('Rapport des Sites et Stations Météorologiques', 0)

    # Organiser par pays
    sites_data_grouped = sites_data.groupby('pays')

    for country, group in sites_data_grouped:
        doc.add_heading(f'{country}', level=1)
        
        for _, row in group.iterrows():
            doc.add_heading(f"Site: {row['site']}", level=2)
            doc.add_paragraph(f"Latitude: {row['latitude']}, Longitude: {row['longitude']}")
            doc.add_paragraph(f"Pays: {row['pays']}")
            doc.add_paragraph(f"Latitude: {row['latitude']}, Longitude: {row['longitude']}")
            doc.add_paragraph(f"Nom de la station: {row['site']}")
            doc.add_paragraph(f"Distance estimée à la station: {round(float(row['distance']), 1)} km")
            doc.add_paragraph("\n")

    doc.save(csv_output_path)

# Fonction principale
def main():
    # Charger les sites
    sites_data = charger_sites('modele_sites.csv')

    # Demander la période
    start_date = simpledialog.askstring("Date de début", "Entrez la date de début (format YYYY-MM-DD):")
    end_date = simpledialog.askstring("Date de fin", "Entrez la date de fin (format YYYY-MM-DD):")
    start = datetime.strptime(start_date, '%Y-%m-%d')
    end = datetime.strptime(end_date, '%Y-%m-%d')

    # Traiter chaque site
    for _, row in sites_data.iterrows():
        print(f"📍 Traitement de {row['site']}...")
        lat, lon = row['latitude'], row['longitude']
        
        # Récupérer les données de vent depuis Meteostat
        wind_meteostat, station = get_meteostat_wind_data(lat, lon, start, end)
        
        if wind_meteostat is not None:
            print(f"📡 Station sélectionnée : {station['name']} (ID: {station['id']})")
            wind_meteostat.to_csv(f"{row['site']}_meteostat_wind.csv")  # Sauvegarder les données dans un fichier CSV
        else:
            print(f"⚠️ Aucune donnée disponible pour {row['site']}")

    # Générer un rapport des sites traités
    generate_report(sites_data, 'rapport_sites.docx')

# Lancer la fonction principale
if __name__ == "__main__":
    main()
