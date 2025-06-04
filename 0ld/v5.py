import os
import time
import pandas as pd
import tkinter as tk
from tkinter import simpledialog, messagebox
from meteostat import Stations, Daily
from datetime import datetime
import requests
from docx import Document

# ------------------------------
# Boîte de dialogue pour date
# ------------------------------
def get_date_input(prompt):
    root = tk.Tk()
    root.withdraw()
    date_str = simpledialog.askstring("Entrée de la date", prompt)
    try:
        return datetime.strptime(date_str, "%Y-%m-%d")
    except:
        messagebox.showerror("Erreur", "Format attendu : YYYY-MM-DD")
        return get_date_input(prompt)

# ------------------------------
# Récupère les données Meteostat
# ------------------------------
def get_meteostat_data(site_name, lat, lon, start, end):
    try:
        stations = Stations().nearby(lat, lon)
        station = stations.fetch(1)
        if station.empty:
            print(f"[⚠️] Aucune station Meteostat pour {site_name}")
            return None, None

        station_id = station.index[0]
        station_info = station.loc[station_id]
        print(f"[📡] Station sélectionnée pour {site_name} : {station_info['name']} (ID: {station_id})")

        data = Daily(station_id, start, end).fetch()
        if data.empty:
            print(f"[⚠️] Aucune donnée disponible pour {site_name} entre {start} et {end}")
            return None, station_info

        print(f"[INFO] Colonnes disponibles dans les données Meteostat : {data.columns.tolist()}")
        data.reset_index(inplace=True)

        df = data[["time", "wspd", "wpgt", "wdir"]].copy()
        df.rename(columns={
            "time": "date",
            "wspd": "wind_speed",
            "wpgt": "wind_gust",
            "wdir": "wind_direction"
        }, inplace=True)
        return df, station_info

    except Exception as e:
        print(f"[ERROR] Problème lors de la récupération des données Meteostat pour {site_name} : {e}")
        return None, None

# ------------------------------
# Appelle l'API OpenMeteo avec retry
# ------------------------------
def get_openmeteo_data(site_name, lat, lon, start, end):
    start_str = start.strftime("%Y-%m-%d")
    end_str = end.strftime("%Y-%m-%d")
    url = (
        f"https://archive-api.open-meteo.com/v1/archive"
        f"?latitude={lat}&longitude={lon}"
        f"&start_date={start_str}&end_date={end_str}"
        f"&daily=windspeed_10m_max,windspeed_10m_mean,winddirection_10m_dominant"
        f"&timezone=auto"
    )

    for attempt in range(3):
        try:
            print(f"[INFO] Requête OpenMeteo : {url}")
            response = requests.get(url)
            if response.status_code == 200:
                json_data = response.json()
                df = pd.DataFrame({
                    "date": json_data["daily"]["time"],
                    "wind_speed_max": json_data["daily"]["windspeed_10m_max"],
                    "wind_speed_mean": json_data["daily"]["windspeed_10m_mean"],
                    "wind_direction_dominant": json_data["daily"]["winddirection_10m_dominant"]
                })
                return df
            elif response.status_code == 429:
                print(f"[ERROR] OpenMeteo - Erreur de requête: 429")
                print("[DEBUG] Réponse API:", response.text)
                time.sleep(10)  # Attente avant retry
            else:
                print(f"[ERROR] OpenMeteo - Erreur de requête: {response.status_code}")
                return None
        except Exception as e:
            print(f"[ERROR] OpenMeteo - Exception: {e}")
            return None
    return None

# ------------------------------
# Fonction principale
# ------------------------------
def main():
    df_sites = pd.read_csv("modele_sites.csv", delimiter=";", encoding="utf-8")
    start = get_date_input("Entrez la date de début (YYYY-MM-DD) :")
    end = get_date_input("Entrez la date de fin (YYYY-MM-DD) :")

    # Création du dossier principal
    os.makedirs("data", exist_ok=True)

    # Création du rapport Word
    doc = Document()
    doc.add_heading("Résumé des données météo", level=1)

    for _, row in df_sites.iterrows():
        site_name = row["Nom"]
        lat = float(row["Latitude"])
        lon = float(row["Longitude"])
        print(f"\n📍 Traitement de {site_name}...")

        output_dir = os.path.join("data", site_name.replace(" ", "_"))
        os.makedirs(output_dir, exist_ok=True)

        # ---- Meteostat
        meteostat_df, station_info = get_meteostat_data(site_name, lat, lon, start, end)
        if meteostat_df is not None:
            meteostat_csv = os.path.join(output_dir, f"meteostat_{site_name.replace(' ', '_')}.csv")
            meteostat_df.to_csv(meteostat_csv, index=False)
            print(f"[✅] Données Meteostat sauvegardées dans : {meteostat_csv}")
        else:
            print(f"[⚠️] Aucune donnée Meteostat pour {site_name}")

        # ---- OpenMeteo
        openmeteo_df = get_openmeteo_data(site_name, lat, lon, start, end)
        time.sleep(25)  # Toujours attendre avant de passer au site suivant

        if openmeteo_df is not None:
            openmeteo_csv = os.path.join(output_dir, f"openmeteo_{site_name.replace(' ', '_')}.csv")
            openmeteo_df.to_csv(openmeteo_csv, index=False)
            print(f"[✅] Données OpenMeteo sauvegardées dans : {openmeteo_csv}")
        else:
            print(f"[⚠️] Aucune donnée OpenMeteo pour {site_name}")

        # ---- Ajouter au rapport Word
        doc.add_heading(site_name, level=2)
        doc.add_paragraph(f"Latitude : {lat}")
        doc.add_paragraph(f"Longitude : {lon}")
        doc.add_paragraph(f"Période étudiée : {start.date()} → {end.date()}")
        if station_info is not None:
            doc.add_paragraph(f"Station Meteostat : {station_info['name']} (ID: {station_info.name})")
        else:
            doc.add_paragraph("Station Meteostat : non trouvée")

        doc.add_paragraph(" ")

    rapport_path = os.path.join("data", "Résumé_Météo.docx")
    doc.save(rapport_path)
    print(f"\n📄 Rapport Word généré dans : {rapport_path}")

if __name__ == "__main__":
    main()
