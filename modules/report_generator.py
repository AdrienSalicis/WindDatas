import os
from docx import Document
from docx.shared import Inches
import pandas as pd
from docx2pdf import convert
from datetime import datetime

def generate_site_report(site_data, output_folder, template_path='template.docx'):
    """
    Génère un rapport DOCX et PDF pour un site.
    Utilise un template Word pour style pro.
    """

    site_ref = f"{site_data['reference']}_{site_data['name']}"
    base_dir = os.path.join(output_folder, site_ref)
    report_dir = os.path.join(base_dir, "report")
    figures_dir = os.path.join(base_dir, "figures")
    tables_dir = os.path.join(base_dir, "tables")

    os.makedirs(report_dir, exist_ok=True)

    # Charger le template
    doc = Document(template_path)

    # --- Page de garde
    doc.add_heading(f"Rapport technique – {site_data['name']}", 0)
    doc.add_paragraph(f"Pays : {site_data.get('country','N/A')}")
    doc.add_paragraph(f"Période d'étude : {site_data.get('start','N/A')} - {site_data.get('end','N/A')}")
    doc.add_paragraph(f"Date de génération : {datetime.now().strftime('%d/%m/%Y')}")

    doc.add_page_break()

    # --- Table des matières
    doc.add_heading('Table des matières', level=1)
    doc.add_paragraph('1. Informations générales')
    doc.add_paragraph('2. Métadonnées des stations')
    doc.add_paragraph('3. Résultats statistiques')
    doc.add_paragraph('4. Graphiques générés')
    doc.add_page_break()

    # --- 1. Informations générales
    doc.add_heading('1. Informations générales', level=1)
    doc.add_paragraph(f"Nom du site : {site_data['name']}")
    doc.add_paragraph(f"Pays : {site_data.get('country','N/A')}")
    doc.add_paragraph(f"Latitude : {site_data.get('latitude','N/A')}")
    doc.add_paragraph(f"Longitude : {site_data.get('longitude','N/A')}")
    doc.add_paragraph(f"Période d'étude : {site_data.get('start','N/A')} - {site_data.get('end','N/A')}")

    # --- 2. Métadonnées des stations
    doc.add_heading('2. Métadonnées des stations', level=1)
    for key in ["meteostat1", "meteostat2", "noaa1", "noaa2"]:
        station = site_data.get(key)
        if station:
            doc.add_heading(f"{key.upper()}", level=2)
            for k, v in station.items():
                doc.add_paragraph(f"{k} : {v}")

    doc.add_page_break()

    # --- 3. Résultats statistiques
    doc.add_heading('3. Résultats statistiques', level=1)
    stats_path = os.path.join(tables_dir, f"stats_descriptives_{site_ref}.csv")
    if os.path.exists(stats_path):
        df_stats = pd.read_csv(stats_path)
        doc.add_paragraph('Tableau des statistiques descriptives :')
        table = doc.add_table(rows=(len(df_stats)+1), cols=len(df_stats.columns))
        table.style = 'Table Grid'
        for j, col in enumerate(df_stats.columns):
            table.cell(0, j).text = col
        for i, row in df_stats.iterrows():
            for j, col in enumerate(df_stats.columns):
                table.cell(i+1, j).text = str(row[col])
    else:
        doc.add_paragraph("Aucun fichier de statistiques trouvé.")

    doc.add_page_break()

    # --- 4. Graphiques générés
    doc.add_heading('4. Graphiques générés', level=1)
    figures_list = sorted([f for f in os.listdir(figures_dir) if f.endswith('.png')])
    for fig in figures_list:
        doc.add_paragraph(fig.replace("_", " ").split(".png")[0])
        try:
            doc.add_picture(os.path.join(figures_dir, fig), width=Inches(5.5))
        except Exception as e:
            doc.add_paragraph(f"[Erreur lors de l'insertion de l'image : {e}]")
        doc.add_paragraph("")

    # --- Sauvegarde DOCX
    output_docx = os.path.join(report_dir, f"fiche_{site_ref}.docx")
    doc.save(output_docx)
    print(f"[✅] Rapport DOCX généré : {output_docx}")

    # --- Conversion automatique en PDF
    try:
        output_pdf = os.path.join(report_dir, f"fiche_{site_ref}.pdf")
        convert(output_docx, output_pdf)
        print(f"[✅] PDF généré automatiquement : {output_pdf}")
    except Exception as e:
        print(f"[⚠️] Erreur conversion PDF : {e}")

    return output_docx
