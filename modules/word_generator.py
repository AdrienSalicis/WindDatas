
import os
from docx import Document
from docx.shared import Pt

def create_word_report_by_country(sites_data, output_path):
    document = Document()
    document.add_heading('🌍 Rapport de Données Météo par Pays', level=0)

    # Regrouper les sites par pays
    grouped = {}
    for site in sites_data:
        grouped.setdefault(site['country'], []).append(site)

    for country, sites in grouped.items():
        document.add_heading(f'📍 {country}', level=1)
        for site in sites:
            name = site['name']
            lat = site['latitude']
            lon = site['longitude']
            start = site['start']
            end = site['end']

            ms1 = site['meteostat1']
            ms2 = site['meteostat2']

            document.add_heading(f"🛰️ Site : {name}", level=2)
            document.add_paragraph(f"Coordonnées GPS : {lat}, {lon}")
            document.add_paragraph(f"Période d’étude : {start} → {end}")
            document.add_paragraph("Variables étudiées : Vitesse et direction du vent, neige")

            document.add_paragraph(f"Meteostat Station 1 : {ms1['id']} – {ms1['name']} ({ms1['distance_km']:.2f} km)")
            document.add_paragraph(f"Meteostat Station 2 : {ms2['id']} – {ms2['name']} ({ms2['distance_km']:.2f} km)")
            document.add_paragraph("OpenMeteo : Données calculées via API")

            document.add_paragraph("📁 Fichiers générés :")
            for f in site['files']:
                p = document.add_paragraph(f"• {f}")
                p.paragraph_format.left_indent = Pt(20)

    if not os.path.exists(os.path.dirname(output_path)):
        os.makedirs(os.path.dirname(output_path))

    document.save(output_path)
    return output_path
